#!/usr/bin/env python3
"""
Removes Course (Year 9) — Teacher Lesson Plans generator.

Produces a SINGLE Word document containing one 40-minute lesson plan per
lesson, each starting on a fresh page. Content is derived faithfully from
the live course slides in js/slides-gcse.js.

Usage:
    python3 build_lesson_plans.py [out.docx]
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x9b, 0x18, 0x44)      # Removes maroon
INK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)

# ---------------------------------------------------------------------------
# LESSON DATA — one dict per lesson, derived from the live slides.
# ---------------------------------------------------------------------------
LESSONS = [
    # ===================================================================
    {
        "id": 101,
        "unit": "Unit 1: Understanding AI",
        "ailit": "Engage with AI",
        "title": "What Is AI?",
        "big_idea": "AI is pattern recognition at massive scale — not thinking, "
                    "not magic, not consciousness. Understanding that one fact is "
                    "what gives students power over the technology rather than the "
                    "other way around.",
        "objectives": [
            "Define AI as statistical pattern recognition learned from data — not thinking or understanding.",
            "Explain the difference between traditional rule-based software and machine learning.",
            "Describe how a machine 'learns' (training set vs test set) and what overfitting / shortcut learning means.",
            "Identify which everyday tools genuinely use AI and which just follow fixed rules.",
        ],
        "vocab": [
            ("Machine learning", "A system that learns its own rules from examples rather than being given them."),
            ("Training set / test set", "Data used to teach the model vs unseen data used to check it really learned."),
            ("Overfitting", "Memorising the training examples but failing on anything new."),
            ("Shortcut learning", "Learning the wrong signal — e.g. the weather in a photo, not the tank in it (Geirhos 2020)."),
            ("Adversarial example", "A tiny, human-invisible pixel change that flips an AI's answer (panda → gibbon)."),
        ],
        "resources": [
            ("Course slides — Lesson 101 'What Is AI?'", "Removes course, Unit 1"),
            ("Video: 'But what is a neural network?' — 3Blue1Brown (first 6 min)", "https://www.youtube.com/watch?v=aircAruvnKk"),
            ("Gemini (school Google account)", "https://gemini.google.com"),
            ("Source: Geirhos et al. (2020), Shortcut Learning in Deep Neural Networks", "https://www.nature.com/articles/s42256-020-00257-z"),
        ],
        "sequence": [
            ("Starter — Hook", "5 min",
             "Display the hook stats: ChatGPT passed 800M weekly users (OpenAI, Oct 2025) — the fastest "
             "adoption of any technology in history — yet only ~23% of young people can explain what AI "
             "actually does. Pose the question on the board: 'You probably use it every week. Could you "
             "explain it to a 10-year-old?' Take 3–4 verbal answers, don't correct them yet."),
            ("Watch — Neural networks", "6 min",
             "Play the first 6 minutes of the 3Blue1Brown video. Set a focus task before pressing play: "
             "'Watch for what \"learning\" actually means.' (Answer: adjusting numbers/weights to reduce error.)"),
            ("Main teaching — What AI really is", "10 min",
             "Teach the three concept slides: (1) AI = prediction from patterns, no knowing/understanding; "
             "(2) Machine learning vs the old way — recipe vs training a chef; spam-filter example; "
             "(3) How a machine learns unplugged — the 'cat photos' analogy, training vs test set, the "
             "'tank/weather' shortcut-learning story. Emphasise: confident ≠ correct."),
            ("Activity — Gemini classify + footprint audit", "12 min",
             "Students open Gemini. Task A (paired): give Gemini six everyday tools (TikTok next video, a "
             "calculator, Netflix recommendations, an alarm clock, a spam filter, face unlock) and ask it "
             "to sort each into 'learned from data' vs 'fixed rule a human wrote', with a one-line reason. "
             "Students then critique Gemini's answers against the lesson. Task B (individual, in notes box): "
             "the 24-hour AI footprint audit — list every AI interaction since yesterday, star the top 3 "
             "that most shape what they see or do."),
            ("Discussion — The picture that broke the AI", "5 min",
             "Present the 2015 adversarial panda → gibbon example. Discuss: does the AI 'see' the way we "
             "do? Draw out that it computes statistical patterns over pixels, which is why it is powerful "
             "AND brittle — and why high-stakes image AI runs with human oversight."),
            ("Plenary — Exit ticket", "2 min",
             "On a slip / in the notes box: 'Finish this sentence in your own words: AI is ______, not ______.' "
             "Target answer shape: 'pattern recognition / statistics, not thinking / understanding.'"),
        ],
        "discussion": [
            "If AI doesn't 'understand' anything, why does it feel like it does when you chat to it?",
            "The calculator and the alarm clock are predictable but brittle; the AI tools are powerful but can be confidently wrong. Which would you rather trust with something important, and why?",
            "Shortcut learning made an AI detect the weather instead of the tank. Where could that kind of hidden mistake be dangerous in the real world?",
        ],
        "class_task": (
            "Gemini fact-check challenge: in pairs, ask Gemini 'Explain how AI learns, in 3 sentences a "
            "13-year-old would understand.' Then mark its answer against today's lesson — is anything "
            "oversimplified, missing, or wrong? Write a better version. (Reinforces: even the tool you're "
            "learning about can be fluently imperfect.)"
        ),
        "differentiation": (
            "Support: provide the six classify items pre-printed so students sort on paper before using "
            "Gemini. Stretch: ask students to find a seventh everyday tool and justify which category it "
            "belongs in — and a borderline case that's genuinely hard to classify."
        ),
        "assessment": (
            "Exit-ticket sentence (AfL); quality of reasoning in the Gemini critique; the in-course quiz "
            "('Which is the most accurate description of what AI does?' — answer: finds patterns in data "
            "to make predictions)."
        ),
        "notes": (
            "The 'tank/weather' story is partly folklore — say so. The phenomenon (shortcut learning) is "
            "real and peer-reviewed (Geirhos 2020). Keep the lesson's honesty about this; it models good "
            "source habits."
        ),
    },
    # ===================================================================
    {
        "id": 102,
        "unit": "Unit 1: Understanding AI",
        "ailit": "Engage with AI",
        "title": "How Chatbots Work",
        "big_idea": "Every chatbot is a 'transformer' that does one thing — predict "
                    "the next token, over and over. There is no understanding in the "
                    "pipeline, which is exactly why confident output can be completely "
                    "wrong (hallucination).",
        "objectives": [
            "Explain that chatbots generate text by predicting the next token, one at a time.",
            "Describe what a token is and why AI 'sees' chunks/numbers, not letters (the strawberry problem).",
            "Explain why hallucinations happen and where they are most likely (recent events, exact numbers, citations).",
            "Judge when a chatbot answer is low-risk vs high-risk to trust without checking.",
        ],
        "vocab": [
            ("Token", "A word-sized chunk of text (~0.75 words). The model only ever sees tokens as numbers."),
            ("Next-token prediction", "Picking the most likely next chunk from a fixed vocabulary, then repeating."),
            ("Transformer", "The neural-network design behind every modern chatbot (Vaswani et al. 2017)."),
            ("Hallucination", "A confident, plausible-sounding answer that is simply false."),
            ("Knowledge cut-off", "The date training stopped; anything after may be missing or wrong."),
            ("Temperature", "A setting controlling how often the model picks a less-likely word (more creative / more risky)."),
        ],
        "resources": [
            ("Course slides — Lesson 102 'How Chatbots Work'", "Removes course, Unit 1"),
            ("Video: 'But what is a GPT?' — 3Blue1Brown (first 7 min)", "https://www.youtube.com/watch?v=wjZofJX0v4M"),
            ("Gemini (school Google account)", "https://gemini.google.com"),
            ("Source: Vaswani et al. (2017), 'Attention Is All You Need'", "https://arxiv.org/abs/1706.03762"),
            ("Source: Stanford HAI (2024) on legal-AI hallucination rates", "https://hai.stanford.edu/news/ai-legal-research-tools-matter-hallucinations"),
        ],
        "sequence": [
            ("Starter — Hook", "5 min",
             "Hook stats: the transformer paper (2017) has 140,000+ citations; every chatbot works by "
             "predicting one word at a time; zero understanding is involved. Write on the board: "
             "'Confident output ≠ correct output.' Ask students to predict the next word in: 'The capital "
             "of France is ___' (easy) then 'The capital of Australia is ___' (many will say Sydney — wrong, "
             "it's Canberra). That misfire IS the lesson."),
            ("Watch — What is a GPT?", "6 min",
             "Play the first 6–7 minutes of the 3Blue1Brown GPT video (tokens + next-word prediction). "
             "Focus task: 'It's maths, not magic — what is the model actually choosing each step?'"),
            ("Main teaching — Tokens, prediction, hallucination", "10 min",
             "Teach: tokens and prediction (the whole trick is next-token prediction); the strawberry "
             "problem (the model sees /Str/aw/berry/ as numbers, so it can't count letters); training data "
             "(no fact-checker in the pipeline); why hallucinations happen (no 'I don't know' default — it "
             "must output something plausible). Cite Stanford 2024: 17–33% hallucination even on "
             "specialist legal AI."),
            ("Activity — Break a chatbot in 5 minutes (Gemini)", "13 min",
             "Students open Gemini and run four tests, recording results in the notes box: "
             "(1) count the t's in 'The quick brown fox jumps over the lazy dog'; "
             "(2) '23 × 47, show the steps' then check on a calculator; "
             "(3) 'Who won the most recent FIFA World Cup?' and compare to reality; "
             "(4) 'Summarise the key findings of Patel et al. (2023) on teenage sleep' — a paper that may "
             "not exist (watch it fabricate). For each: was it confident? Was it right?"),
            ("Discussion — Sort the risk", "4 min",
             "Using the classify task: which jobs are safe to trust without checking (rewrite my paragraph, "
             "brainstorm story ideas, summarise text I pasted) vs high-risk (exact population, cite three "
             "studies, count a specific letter)? Draw out the rule: language/idea tasks low-risk; facts, "
             "numbers, citations, character-level tasks high-risk."),
            ("Plenary — Exit ticket", "2 min",
             "'Name one task you would trust a chatbot with, and one you would always verify — and say why "
             "in one sentence each.'"),
        ],
        "discussion": [
            "If a chatbot has no fact-checker inside it, why does it so often happen to be right?",
            "Why is a fabricated, realistic-looking citation more dangerous than an obvious mistake?",
            "Knowing the model predicts 'plausible' not 'true', how should that change the way you phrase a question or use the answer?",
        ],
        "class_task": (
            "Hallucination hunt (Gemini): each pair tries to get Gemini to confidently state something "
            "false — a fake book, a made-up statistic, a non-existent study — and screenshots/records it. "
            "Class shares the best examples and labels WHERE on the 'edge of knowledge' each one sat "
            "(recent? niche? exact number? specific person?)."
        ),
        "differentiation": (
            "Support: run the four 'break a chatbot' tests as a teacher-led demo on the board, students "
            "predict each outcome first. Stretch: students investigate the 'temperature' idea — ask Gemini "
            "the same creative prompt three times and explain why answers vary."
        ),
        "assessment": (
            "Exit ticket (trust/verify judgement); recorded results from the four tests with a written "
            "'what this tells me' line; in-course quiz (London-as-capital-of-Australia → plausible "
            "prediction, not a fact)."
        ),
        "notes": (
            "Reasoning models now usually get 'strawberry' right because they spell it out first — say so, "
            "then show the weakness still exists with a longer/made-up word. Keeps the lesson current and "
            "honest rather than claiming the flaw is gone."
        ),
    },
    # ===================================================================
    {
        "id": 103,
        "unit": "Unit 1: Understanding AI",
        "ailit": "Engage with AI",
        "title": "Types of AI",
        "big_idea": "Every AI in existence today is narrow — superhuman at one task, "
                    "infant-level at everything else. General AI (AGI) does not exist. "
                    "The honest test for any 'AI will do X' claim is whether X is "
                    "narrow and well-defined or open-ended and cross-domain.",
        "objectives": [
            "Distinguish narrow AI (ANI) from general AI (AGI) and explain that all current AI is narrow.",
            "Name the main types of AI (generative, classifiers, recommenders, vision, speech, reinforcement learning).",
            "Match real-world tools to their underlying AI type(s).",
            "Use the narrow-vs-general test to evaluate hype claims about what AI will replace.",
        ],
        "vocab": [
            ("Narrow AI (ANI)", "Built for one task; superhuman inside it, useless outside it. Every product today."),
            ("General AI (AGI)", "Could flex across any task a human can. Does not exist yet."),
            ("Generative AI", "Creates new content — text, images, music (ChatGPT, Midjourney, Suno)."),
            ("Classifier", "Sorts things into categories — spam/not spam, tumour/no tumour."),
            ("Recommendation engine", "Predicts what you'll want next (TikTok, Spotify, Netflix)."),
            ("Reinforcement learning", "Learns by trial-and-error against a reward (AlphaZero, RLHF)."),
            ("Transfer learning", "Applying a skill learned in one task to a brand-new one — easy for a child, beyond current AI."),
        ],
        "resources": [
            ("Course slides — Lesson 103 'Types of AI'", "Removes course, Unit 1"),
            ("Gemini (school Google account)", "https://gemini.google.com"),
            ("NotebookLM (for the AlphaFold case study)", "https://notebooklm.google.com"),
            ("Source: Silver et al. (DeepMind, Science 2018) — AlphaZero", "https://www.science.org/doi/10.1126/science.aar6404"),
            ("Source: Nobel Prize in Chemistry 2024 — AlphaFold", "https://www.nobelprize.org/prizes/chemistry/2024/press-release/"),
        ],
        "sequence": [
            ("Starter — Hook", "5 min",
             "Hook: in 2017 AlphaZero taught itself chess from scratch and beat the best engine on Earth "
             "in 4 hours of self-play — yet it can't make tea, write an essay, or recognise a face. Ask: "
             "'Is AlphaZero more intelligent than you?' Take a quick show-of-hands and two reasons each way."),
            ("Main teaching — Narrow vs general", "10 min",
             "Teach narrow AI (ANI) vs general AI (AGI): every system today is narrow. AGI timelines range "
             "from 5 years (Altman, 2024) to never (LeCun, 2024). The trap: a chatbot FEELS general because "
             "text covers so much, but it's still one narrow domain. Introduce Moravec's Paradox — easy for "
             "a 5-year-old is hard for AI, and vice versa."),
            ("Activity — Narrow or general? + match the type", "12 min",
             "Task A (whole-class, Gemini optional): the 'Narrow or General?' sort — for each capability "
             "(beat a Go champion / flag a tumour / ride a bike then skateboard the next day / write an "
             "essay / read a quiet friend's mood / recommend a video) decide: can today's narrow AI do it, "
             "or would it need AGI? Task B (paired): match the tool to the type — ChatGPT, TikTok feed, "
             "Spotify, Siri/Alexa, Google Photos face grouping, Gmail spam. Note which tools stack several "
             "types."),
            ("Case study — AlphaFold (NotebookLM)", "8 min",
             "Brief students on AlphaFold: a narrow AI that solved 50-year protein-folding problem, mapped "
             "200M+ proteins, used by 2M+ researchers, won the 2024 Nobel Prize in Chemistry. Optional: "
             "load the DeepMind/Nobel sources into NotebookLM and have students ask it 'Why does AlphaFold "
             "matter if it's only narrow AI?' Draw out: narrow ≠ lesser."),
            ("Discussion — Think & discuss", "3 min",
             "Quick-fire: 'If an AI beats every human at chess, does that make it more intelligent than us? "
             "What would \"generally\" intelligent even mean?'"),
            ("Plenary — Exit ticket", "2 min",
             "'A company claims its AI will replace teachers. Using today's lesson, give ONE reason that "
             "claim is probably overblown.' (Target: teaching is open-ended/cross-domain, not narrow.)"),
        ],
        "discussion": [
            "If an AI can beat any human at chess, does that make it more intelligent than us — or just better at chess?",
            "What would it actually take for an AI to be 'generally' intelligent, and why is that so hard to build?",
            "AlphaFold (narrow) won a Nobel Prize while chatbots grab the headlines. Does the most useful AI get the most attention?",
        ],
        "class_task": (
            "Hype audit (Gemini): students paste a real 'AI will replace [job]' headline (or one you "
            "provide) into Gemini and ask it to break the job into narrow vs open-ended tasks. They then "
            "decide, with reasons, which parts current narrow AI could do and which would need AGI — and "
            "write a one-paragraph verdict on the headline."
        ),
        "differentiation": (
            "Support: pre-sort three of the 'narrow or general?' items together as a class before students "
            "do the rest. Stretch: students argue both sides of an AGI timeline (Altman vs LeCun) and "
            "state which they find more convincing and why."
        ),
        "assessment": (
            "Exit-ticket reasoning; accuracy of the tool-to-type matching; the in-course quiz (chess AI "
            "that can't write or recognise faces → all current AI is narrow)."
        ),
        "notes": (
            "Keep model names current when students ask: as of mid-2026 the leading general-purpose "
            "chatbots are GPT-5.x, Claude (Opus 4.8 / Sonnet 4.6) and Gemini 3.x. The teaching point is "
            "unchanged: however capable, they are still narrow."
        ),
    },
]

# ---------------------------------------------------------------------------
# Rendering helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, text, size=14, color=BRAND, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(doc, text, size=10.5, color=INK, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = INK
    return p


def render_lesson(doc, L, first=False):
    if not first:
        add_page_break(doc)

    # Title block
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Lesson {L['id']} · {L['unit']}")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    r.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(L['title'])
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"Duration: 40 minutes   ·   AI Literacy domain: {L['ailit']}")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    # Big idea callout
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'F6E8EE')
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    rr = cp.add_run("Big idea  ")
    rr.bold = True
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = BRAND
    rr2 = cp.add_run(L['big_idea'])
    rr2.font.size = Pt(10.5)
    rr2.font.color.rgb = INK

    # Learning objectives
    heading(doc, "Learning objectives", size=13)
    body(doc, "By the end of the lesson, students will be able to:", bold=True, space_after=2)
    for o in L['objectives']:
        bullet(doc, o)

    # Key vocabulary
    heading(doc, "Key vocabulary", size=13)
    vt = doc.add_table(rows=0, cols=2)
    vt.style = 'Table Grid'
    vt.columns[0].width = Inches(1.8)
    vt.columns[1].width = Inches(4.7)
    for term, definition in L['vocab']:
        row = vt.add_row().cells
        rp = row[0].paragraphs[0]
        rr = rp.add_run(term)
        rr.bold = True
        rr.font.size = Pt(10)
        rr.font.color.rgb = BRAND
        dp = row[1].paragraphs[0]
        dr = dp.add_run(definition)
        dr.font.size = Pt(10)
        dr.font.color.rgb = INK

    # Resources
    heading(doc, "Resources needed", size=13)
    for label, link in L['resources']:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
        if link and link.startswith('http'):
            r2 = p.add_run(f"  —  {link}")
            r2.font.size = Pt(9)
            r2.font.color.rgb = GREY
        elif link:
            r2 = p.add_run(f"  ({link})")
            r2.font.size = Pt(9)
            r2.font.color.rgb = GREY

    # Lesson sequence
    heading(doc, "Lesson sequence (40 minutes)", size=13)
    st = doc.add_table(rows=1, cols=3)
    st.style = 'Table Grid'
    st.columns[0].width = Inches(1.9)
    st.columns[1].width = Inches(0.7)
    st.columns[2].width = Inches(3.9)
    hdr = st.rows[0].cells
    for i, h in enumerate(("Phase", "Time", "What happens")):
        set_cell_bg(hdr[i], '9B1844')
        hp = hdr[i].paragraphs[0]
        hr = hp.add_run(h)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for phase, time, what in L['sequence']:
        row = st.add_row().cells
        pp = row[0].paragraphs[0]
        pr = pp.add_run(phase)
        pr.bold = True
        pr.font.size = Pt(10)
        pr.font.color.rgb = INK
        tp = row[1].paragraphs[0]
        tr = tp.add_run(time)
        tr.font.size = Pt(10)
        tr.font.color.rgb = GREY
        wp = row[2].paragraphs[0]
        wr = wp.add_run(what)
        wr.font.size = Pt(10)
        wr.font.color.rgb = INK

    # Discussion questions
    heading(doc, "Discussion questions", size=13)
    for i, q in enumerate(L['discussion'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i}.  ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = BRAND
        r2 = p.add_run(q)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = INK

    # Class task with AI tools
    heading(doc, "Class task — using Gemini / NotebookLM", size=13)
    body(doc, L['class_task'])

    # Differentiation
    heading(doc, "Differentiation", size=13)
    body(doc, L['differentiation'])

    # Assessment
    heading(doc, "Assessment for learning", size=13)
    body(doc, L['assessment'])

    # Teacher notes
    heading(doc, "Teacher notes", size=13)
    body(doc, L['notes'])


# ---------------------------------------------------------------------------
def build(out_path):
    doc = Document()

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # Cover / intro
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Digital Innovations — Removes Course (Year 9)")
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Teacher Lesson Plans")
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = BRAND

    body(doc,
         "One 40-minute lesson plan per lesson. Each plan begins on a new page and is derived "
         "directly from the live course content. Plans build in discussion questions and class "
         "tasks that use Gemini or NotebookLM during the lesson. Aligned to the OECD/EU AI "
         "Literacy Framework (Engage · Create · Manage · Shape).",
         size=11)
    body(doc,
         "This draft contains the first three lessons (101–103) for review. The remaining lessons "
         "and an overarching Scheme of Work will follow once the format is approved.",
         size=10.5, color=GREY)

    for i, L in enumerate(LESSONS):
        render_lesson(doc, L, first=(i == 0))

    doc.save(out_path)
    print(f"Saved {out_path} with {len(LESSONS)} lesson plan(s).")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "Removes-Lesson-Plans.docx"
    build(out)
