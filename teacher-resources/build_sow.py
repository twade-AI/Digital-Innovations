#!/usr/bin/env python3
"""
Build the overarching Scheme of Work (SOW) for the Removes course as a Word
document. One row per lesson, grouped by unit, in teaching order. Columns:
lesson, title, learning objectives, key tasks, resources.

Reuses the lesson data assembled by build_full.py.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

import build_lesson_plans as base
from build_full import collect

BRAND = base.BRAND
GREY = base.GREY
INK = base.INK
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def cell_text(cell, text, size=8.5, bold=False, color=INK, after=2):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color
    return p


def cell_bullets(cell, items, size=8.5):
    cell.paragraphs[0].text = ""
    first = True
    for it in items:
        if first:
            p = cell.paragraphs[0]; first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.08)
        r = p.add_run("• " + it)
        r.font.size = Pt(size); r.font.color.rgb = INK


def key_tasks_for(L):
    """Pull the hands-on phases (activity/discussion/plenary) as 'key tasks'."""
    tasks = []
    for phase, time, what in L["sequence"]:
        pl = phase.lower()
        if any(k in pl for k in ("activity", "case", "discussion", "challenge", "plenary", "watch")):
            # short label = phase; keep it terse
            tasks.append(f"{phase} ({time})")
    return tasks


def build(out_path):
    lessons = collect()
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Digital Innovations — Removes Course (Year 9)")
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = GREY
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Scheme of Work")
    r.font.size = Pt(22); r.bold = True; r.font.color.rgb = BRAND
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("7 units · 40 lessons · 40 minutes each · aligned to the OECD/EU AI Literacy Framework "
                  "(Engage · Create · Manage · Shape). Lessons are listed in teaching order.")
    r.font.size = Pt(9.5); r.font.color.rgb = INK

    # Table
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Inches(0.55), Inches(1.7), Inches(3.2), Inches(1.9), Inches(2.85)]
    headers = ["Lesson", "Title", "Learning objectives", "Key tasks", "Resources"]
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        base.set_cell_bg(hdr[i], '9B1844')
        cell_text(hdr[i], h, size=9, bold=True, color=WHITE, after=0)

    current_unit = None
    for L in lessons:
        # Unit banner row (spanning)
        if L["unit"] != current_unit:
            current_unit = L["unit"]
            brow = tbl.add_row()
            a = brow.cells[0]
            b = brow.cells[-1]
            merged = a.merge(b)
            base.set_cell_bg(merged, 'F0D9E1')
            cell_text(merged, f"{current_unit}   ·   {L['ailit']}", size=9.5, bold=True,
                      color=BRAND, after=0)

        row = tbl.add_row().cells
        cell_text(row[0], f"L{L['id']}", size=9, bold=True, color=BRAND, after=0)
        cell_text(row[1], L["title"], size=9, bold=True, color=INK, after=0)
        cell_bullets(row[2], L["objectives"], size=8.5)
        cell_bullets(row[3], key_tasks_for(L), size=8.5)
        # resources: label only, links in italic-grey under
        row[4].paragraphs[0].text = ""
        first = True
        for label, link in L["resources"]:
            if first:
                rp = row[4].paragraphs[0]; first = False
            else:
                rp = row[4].add_paragraph()
            rp.paragraph_format.space_after = Pt(1)
            rr = rp.add_run("• " + label)
            rr.font.size = Pt(8); rr.font.color.rgb = INK
            if link and link.startswith("http"):
                lr = rp.add_run(f"  {link}")
                lr.font.size = Pt(7); lr.font.color.rgb = GREY

    # Set column widths
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.save(out_path)
    print(f"Saved {out_path} ({len(lessons)} lessons).")


if __name__ == "__main__":
    build("Removes-Scheme-of-Work.docx")
