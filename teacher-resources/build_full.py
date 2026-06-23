#!/usr/bin/env python3
"""
Assemble ALL 40 Removes lesson plans into a single Word document, in the
real teaching order defined by GCSE_UNITS in js/slides-gcse.js.

Sources:
  - build_lesson_plans.py  -> LESSONS (the 3 worked examples: 101, 102, 103)
                              + render helpers (render_lesson, set_cell_bg, etc.)
  - lessons_batchA/B/C/D.py -> LESSONS_BATCH (the remaining 37)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches

import build_lesson_plans as base
from lessons_batchA import LESSONS_BATCH as A
from lessons_batchB import LESSONS_BATCH as B
from lessons_batchC import LESSONS_BATCH as C
from lessons_batchD import LESSONS_BATCH as D

BRAND = base.BRAND
GREY = base.GREY
INK = base.INK

# Real teaching order (from GCSE_UNITS in js/slides-gcse.js)
TEACHING_ORDER = [
    101, 102, 103, 104, 135, 105,                 # Unit 1
    106, 107, 108, 109, 110, 111, 112,            # Unit 2
    113, 114, 115, 116, 136, 117, 118, 134,       # Unit 3
    119, 120, 121, 122, 123,                      # Unit 4
    124, 125, 137, 126, 127, 128, 129,            # Unit 5
    138, 139, 140,                                # Unit 6
    130, 131, 132, 133,                           # Unit 7
]


def collect():
    by_id = {}
    for L in base.LESSONS + A + B + C + D:
        by_id[L["id"]] = L
    missing = [i for i in TEACHING_ORDER if i not in by_id]
    extra = [i for i in by_id if i not in TEACHING_ORDER]
    if missing:
        raise SystemExit(f"Missing lessons: {missing}")
    if extra:
        raise SystemExit(f"Unexpected lessons: {extra}")
    if len(TEACHING_ORDER) != 40:
        raise SystemExit(f"Expected 40, got {len(TEACHING_ORDER)}")
    return [by_id[i] for i in TEACHING_ORDER]


def build(out_path):
    lessons = collect()
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Digital Innovations — Removes Course (Year 9)")
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Teacher Lesson Plans")
    r.font.size = Pt(26); r.bold = True; r.font.color.rgb = BRAND

    base.body(doc,
        "One 40-minute lesson plan per lesson — all 40 lessons, in teaching order. "
        "Each plan begins on a new page and is derived directly from the live course "
        "content. Plans build in discussion questions and class tasks that use Gemini "
        "or NotebookLM during the lesson. Aligned to the OECD/EU AI Literacy Framework "
        "(Engage · Create · Manage · Shape).", size=11)

    # Contents overview by unit
    base.heading(doc, "Course at a glance", size=14)
    current_unit = None
    for L in lessons:
        if L["unit"] != current_unit:
            current_unit = L["unit"]
            up = doc.add_paragraph()
            up.paragraph_format.space_before = Pt(6)
            up.paragraph_format.space_after = Pt(1)
            ur = up.add_run(f"{current_unit}  ·  {L['ailit']}")
            ur.bold = True; ur.font.size = Pt(10.5); ur.font.color.rgb = BRAND
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(0)
        lp.paragraph_format.left_indent = Inches(0.25)
        lr = lp.add_run(f"L{L['id']}  ·  {L['title']}")
        lr.font.size = Pt(10); lr.font.color.rgb = INK

    # Lesson plans, each on a new page
    for L in lessons:
        base.render_lesson(doc, L, first=False)  # page break before every plan

    doc.save(out_path)
    print(f"Saved {out_path} with {len(lessons)} lesson plans.")


if __name__ == "__main__":
    build("Removes-Lesson-Plans-ALL.docx")
